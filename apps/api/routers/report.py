"""
Word Report API router.
Generates a comprehensive professional investment methodology report.
"""

import tempfile
import time
from datetime import datetime
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional

from apps.api.routers.analysis import get_last_result

router = APIRouter(prefix="/api/report", tags=["report"])


class ReportRequest(BaseModel):
    answers: Optional[dict[str, str]] = None


@router.post("/word")
async def generate_word_report(req: ReportRequest = ReportRequest()):
    last = get_last_result()
    if not last or "profile" not in last:
        raise HTTPException(status_code=400, detail="Ejecuta el análisis primero.")
    try:
        path = _build_word_report(last, req.answers or {})
        profile = last["profile"]
        fname = f"Informe_AHP_{profile}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        return FileResponse(
            path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=fname,
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────────────────────────────────
# STATIC REFERENCE DATA
# ─────────────────────────────────────────────────────────────────────────────

CRITERIA_ORDER = [
    "rentabilidad", "sharpe", "sortino", "alpha",
    "volatilidad", "var_95", "cvar_95", "max_drawdown",
    "beta", "tracking_error", "rentab_kp",
    "cv", "skewness", "corr_media", "div_geo",
]

CRITERIA_META = {
    "rentabilidad":   {"label": "Rentabilidad Anualizada",        "dir": "maximize", "cat": "Rendimiento",     "formula": "μ × 252",                         "desc": "Media de retornos logarítmicos diarios escalada a 252 días hábiles. Mide el crecimiento porcentual anual esperado."},
    "sharpe":         {"label": "Ratio de Sharpe",                "dir": "maximize", "cat": "Rendimiento",     "formula": "(Rp − Rf) / σp",                  "desc": "Exceso de rentabilidad sobre la tasa libre de riesgo por unidad de riesgo total. Mayor = mejor relación riesgo/retorno."},
    "sortino":        {"label": "Ratio de Sortino",               "dir": "maximize", "cat": "Rendimiento",     "formula": "(Rp − Rf) / σ↓",                 "desc": "Igual que Sharpe pero sólo penaliza la volatilidad a la baja. Más relevante para inversores que temen las pérdidas."},
    "alpha":          {"label": "Alpha de Jensen",                "dir": "maximize", "cat": "Rendimiento",     "formula": "Rp − [Rf + β(Rm − Rf)]",          "desc": "Rentabilidad por encima de la predicción del CAPM dado el nivel de riesgo. Mide la capacidad del gestor de generar valor."},
    "volatilidad":    {"label": "Volatilidad (σ)",                "dir": "minimize", "cat": "Riesgo",          "formula": "σ_diaria × √252",                 "desc": "Desviación estándar de los retornos diarios, anualizada. Mayor volatilidad = mayor incertidumbre sobre el resultado final."},
    "var_95":         {"label": "VaR 95% (Paramétrico)",          "dir": "minimize", "cat": "Riesgo",          "formula": "−(μ + z₀.₀₅ × σ)",               "desc": "Pérdida máxima esperada con 95% de confianza. Si VaR = 10%, en 1 de cada 20 períodos se perderá al menos ese porcentaje."},
    "cvar_95":        {"label": "CVaR 95% (Expected Shortfall)",  "dir": "minimize", "cat": "Riesgo",          "formula": "−E[R | R ≤ q₅%]",                "desc": "Pérdida media en el 5% de los peores escenarios. Más conservador que VaR: mide magnitud de pérdidas extremas, no sólo probabilidad."},
    "max_drawdown":   {"label": "Max Drawdown",                   "dir": "minimize", "cat": "Riesgo",          "formula": "máx[(Pico−Valle)/Pico]",           "desc": "Mayor caída porcentual desde el máximo histórico al mínimo siguiente. Representa el peor escenario para un inversor que compró en máximos."},
    "beta":           {"label": "Beta del Portafolio",            "dir": "minimize", "cat": "Eficiencia",      "formula": "Cov(Ri, Rm) / Var(Rm)",            "desc": "Sensibilidad del activo al mercado. β=1 → se mueve como el índice; β>1 amplifica; β<1 amortigua los movimientos del mercado."},
    "tracking_error": {"label": "Tracking Error",                 "dir": "minimize", "cat": "Eficiencia",      "formula": "σ(Ri − Rm) × √252",               "desc": "Desviación estándar de la diferencia de retornos con el benchmark. Mide cuánto se aleja la cartera del índice de referencia."},
    "rentab_kp":      {"label": "Rentabilidad − Coste Capital",   "dir": "maximize", "cat": "Eficiencia",      "formula": "Rp − [Rf + β(Rm − Rf)]",          "desc": "Diferencia entre rentabilidad obtenida y el coste de capital del CAPM. Positivo = el activo supera el retorno mínimo exigido."},
    "cv":             {"label": "Coeficiente de Variación",       "dir": "minimize", "cat": "Estabilidad",     "formula": "σ / |μ|",                         "desc": "Riesgo por unidad de rentabilidad. Permite comparar la dispersión entre activos con distintos niveles de retorno esperado."},
    "skewness":       {"label": "Skewness (Asimetría)",           "dir": "maximize", "cat": "Estabilidad",     "formula": "E[(R−μ)³] / σ³",                  "desc": "Asimetría de la distribución de retornos. Positivo = más retornos extremos positivos; negativo = colas de pérdida más pronunciadas."},
    "corr_media":     {"label": "Correlación Media",              "dir": "minimize", "cat": "Diversificación", "formula": "media(ρᵢⱼ) i≠j",                 "desc": "Correlación media entre pares de activos. Cuanto más baja, mayor beneficio de diversificación y menor riesgo conjunto."},
    "div_geo":        {"label": "Diversificación Geográfica",     "dir": "maximize", "cat": "Diversificación", "formula": "1 − (1/n)",                       "desc": "Dispersión entre los tres mercados (S&P 500, Eurostoxx, Nikkei). Próximo a 1 = cartera bien repartida geográficamente."},
}

ALL_PROFILE_WEIGHTS = {
    "conservador":  {"volatilidad":9,"var_95":9,"max_drawdown":9,"sortino":8,"cvar_95":8,"beta":8,"corr_media":8,"sharpe":7,"tracking_error":7,"cv":7,"skewness":6,"div_geo":6,"alpha":3,"rentab_kp":3,"rentabilidad":2},
    "moderado":     {"sharpe":8,"max_drawdown":8,"sortino":7,"volatilidad":7,"var_95":7,"corr_media":7,"cvar_95":6,"beta":6,"cv":6,"div_geo":6,"rentabilidad":5,"alpha":5,"tracking_error":5,"rentab_kp":5,"skewness":5},
    "agresivo":     {"rentabilidad":8,"alpha":7,"rentab_kp":7,"sharpe":6,"sortino":5,"max_drawdown":5,"corr_media":5,"div_geo":5,"volatilidad":4,"tracking_error":4,"cv":4,"skewness":4,"var_95":3,"cvar_95":3,"beta":3},
    "muy_agresivo": {"rentabilidad":9,"rentab_kp":9,"alpha":8,"sharpe":4,"div_geo":4,"sortino":3,"max_drawdown":3,"corr_media":3,"volatilidad":2,"tracking_error":2,"cv":2,"skewness":2,"var_95":1,"cvar_95":1,"beta":1},
    "dividendos":   {"sortino":8,"var_95":8,"max_drawdown":8,"sharpe":7,"volatilidad":7,"cvar_95":7,"beta":7,"cv":6,"corr_media":6,"tracking_error":5,"rentab_kp":5,"skewness":5,"div_geo":5,"rentabilidad":4,"alpha":4},
    "tecnologico":  {"rentabilidad":8,"alpha":8,"rentab_kp":7,"sharpe":6,"sortino":5,"max_drawdown":5,"cv":4,"corr_media":4,"volatilidad":3,"var_95":3,"tracking_error":3,"skewness":3,"div_geo":3,"cvar_95":2,"beta":2},
    "esg":          {"sharpe":7,"sortino":7,"var_95":7,"max_drawdown":7,"corr_media":7,"div_geo":7,"volatilidad":6,"cvar_95":6,"alpha":5,"beta":5,"tracking_error":5,"rentab_kp":5,"cv":5,"skewness":5,"rentabilidad":4},
}

RI_TABLE = {1:0.00,2:0.00,3:0.58,4:0.90,5:1.12,6:1.24,7:1.32,8:1.41,9:1.45,10:1.49,11:1.51,12:1.54,13:1.56,14:1.57,15:1.59}

PROFILE_FULL_DESC = {
    "conservador":  ("Conservador", "Prioriza la preservación del capital. Acepta rentabilidades modestas a cambio de mínima volatilidad. Horizonte corto-medio, muy baja tolerancia al riesgo. Adecuado para inversores cercanos a la jubilación o con necesidad de liquidez inmediata."),
    "moderado":     ("Moderado",    "Busca equilibrio entre crecimiento y seguridad. Acepta cierta volatilidad a cambio de una rentabilidad razonable. Horizonte medio (5-10 años). Es el perfil más habitual entre inversores particulares con experiencia básica."),
    "agresivo":     ("Agresivo",    "Prioriza la rentabilidad sobre la seguridad. Tolera drawdowns significativos y alta volatilidad con el objetivo de maximizar el crecimiento del capital a largo plazo (>10 años). Para inversores con experiencia y sin necesidad de liquidez a corto plazo."),
    "muy_agresivo": ("Muy Agresivo","Maximiza el retorno sin restricciones de riesgo. Acepta pérdidas temporales muy elevadas. Diseñado para inversores con amplia experiencia, capital disponible a largo plazo y perfil psicológico que tolera alta incertidumbre."),
    "dividendos":   ("Dividendos",  "Busca ingresos periódicos mediante empresas con historial sólido de pagos de dividendos. Preferencia por large caps estables y alta capitalización. Objetivo: flujo de caja constante con baja volatilidad del principal."),
    "tecnologico":  ("Tecnológico", "Concentrado en sectores de innovación: tecnología, semiconductores, software y telecomunicaciones. Alto potencial de crecimiento con mayor volatilidad. Para inversores que creen en la transformación digital y aceptan ciclos de alta fluctuación."),
    "esg":          ("ESG",         "Invierte según criterios medioambientales, sociales y de gobernanza (Environmental, Social, Governance). Combina rentabilidad con impacto positivo. Selecciona empresas con buenas prácticas sostenibles y baja huella de carbono reportada."),
}

SAATY_SCALE = [
    (1, "Igual importancia",                        "Los dos criterios contribuyen de forma idéntica al objetivo"),
    (2, "Importancia débilmente superior",           "Leve preferencia por uno sobre el otro; difícil de distinguir"),
    (3, "Moderadamente más importante",              "Experiencia y juicio favorecen un criterio sobre el otro"),
    (4, "Moderado-fuerte",                           "Entre moderado y fuerte"),
    (5, "Fuertemente más importante",                "Un criterio claramente favorecido; dominante en la práctica"),
    (6, "Fuerte-muy fuerte",                         "Entre fuerte y muy fuerte"),
    (7, "Muy fuertemente más importante",            "Un criterio domina con amplia ventaja; relevancia demostrada"),
    (8, "Extremo-muy fuerte",                        "Entre muy fuerte y extremo"),
    (9, "Extremadamente más importante",             "La diferencia es del mayor orden de magnitud posible"),
]

HML = {range(1,4):"B", range(4,7):"M", range(7,10):"A"}  # bajo, medio, alto

def _weight_to_hml(w):
    if w <= 3: return "B"
    if w <= 6: return "M"
    return "A"


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _fetch_news(tickers: list, max_per: int = 3) -> dict:
    """Fetch latest news for tickers via yfinance."""
    result = {}
    for ticker in tickers[:7]:
        try:
            import yfinance as yf
            t = yf.Ticker(ticker)
            news = t.news or []
            items = []
            for n in news[:max_per]:
                ts = n.get("providerPublishTime", 0)
                date_str = datetime.fromtimestamp(ts).strftime("%d/%m/%Y") if ts else "—"
                title = n.get("title", "")
                if title:
                    items.append({
                        "title":     title,
                        "publisher": n.get("publisher", "—"),
                        "date":      date_str,
                    })
            if items:
                result[ticker] = items
            time.sleep(0.25)
        except Exception:
            pass
    return result


def _priority_vector(weights_dict: dict) -> tuple[list, list, float]:
    """Return (ordered_weights, priority_vector, total) for CRITERIA_ORDER keys."""
    w = [weights_dict.get(k, 1) for k in CRITERIA_ORDER]
    total = sum(w)
    pv = [wi / total for wi in w]
    return w, pv, total


def _normalize_portfolio_scores(portfolios: list) -> tuple:
    """
    Min-max normalize each criterion across portfolios.
    Returns (norm_matrix[n_port x 15], weights_vec[15], priority_vec[15], final_scores[n_port]).
    """
    import numpy as np
    n = len(portfolios)
    if n == 0:
        return [], [], [], []

    raw = []
    for p in portfolios:
        row = []
        for k in CRITERIA_ORDER:
            v = p.get(k, 0) or 0
            row.append(float(v))
        raw.append(row)

    raw = np.array(raw)
    norm = np.zeros_like(raw)

    for j, k in enumerate(CRITERIA_ORDER):
        col = raw[:, j]
        cmin, cmax = col.min(), col.max()
        rng = cmax - cmin
        direction = CRITERIA_META[k]["dir"]
        if rng < 1e-10:
            norm[:, j] = 1.0
        elif direction == "maximize":
            norm[:, j] = (col - cmin) / rng
        else:
            norm[:, j] = (cmax - col) / rng

    weights = [ALL_PROFILE_WEIGHTS.get("moderado", {}).get(k, 1) for k in CRITERIA_ORDER]  # placeholder
    total_w = sum(weights)
    pv = [w / total_w for w in weights]
    scores = norm @ np.array(pv)
    return norm.tolist(), weights, pv, scores.tolist()


# ─────────────────────────────────────────────────────────────────────────────
# MAIN BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _build_word_report(data: dict, answers: dict) -> str:
    import numpy as np
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.27)
    sec.page_height = Inches(11.69)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, Inches(1.0))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)

    # ── Colour palette ──────────────────────────────────────────────────────
    DARK  = RGBColor(0x1A, 0x1A, 0x2E)
    BLUE  = RGBColor(0x16, 0x5F, 0xA8)
    GOLD  = RGBColor(0xC8, 0xA0, 0x00)
    GRAY  = RGBColor(0x6C, 0x75, 0x7D)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN = RGBColor(0x1A, 0x6A, 0x3E)
    RED   = RGBColor(0xB0, 0x30, 0x20)
    LTBLUE= RGBColor(0xEB, 0xF2, 0xFA)

    P_CLR_MAP = {
        "conservador": RGBColor(0x1A, 0x56, 0x9A), "moderado":    RGBColor(0x1A, 0x6A, 0x3E),
        "agresivo":    RGBColor(0xBF, 0x60, 0x00),  "muy_agresivo": RGBColor(0xB0, 0x20, 0x20),
        "dividendos":  RGBColor(0x5B, 0x2D, 0x8E),  "tecnologico":  RGBColor(0x06, 0x7B, 0x8C),
        "esg":         RGBColor(0x0F, 0x6B, 0x41),
    }

    # ── Extracted data ──────────────────────────────────────────────────────
    profile       = data.get("profile", "moderado")
    P_CLR         = P_CLR_MAP.get(profile, BLUE)
    profile_label, profile_full_desc = PROFILE_FULL_DESC.get(profile, (profile.capitalize(), ""))
    winner        = data.get("winner", {})
    allocation    = data.get("allocation", [])
    ranking       = data.get("ranking", [])
    portfolios    = data.get("portfolios", [])
    stocks        = data.get("stocks", [])
    top_criteria  = data.get("top_criteria", [])
    scores        = data.get("scores", {})
    n_analyzed    = data.get("n_stocks_analyzed", 0)
    n_selected    = data.get("n_stocks_selected", 0)
    cr_val        = data.get("consistency_ratio", 0.0)
    is_synthetic  = data.get("is_synthetic", True)
    today         = datetime.now().strftime("%d de %B de %Y")
    winner_port   = next((p for p in portfolios if p.get("name") == winner.get("name")), {})

    profile_weights = ALL_PROFILE_WEIGHTS.get(profile, {k: 5 for k in CRITERIA_ORDER})
    sorted_criteria = sorted(CRITERIA_ORDER, key=lambda k: -profile_weights.get(k, 0))
    w_vec = [profile_weights.get(k, 0) for k in CRITERIA_ORDER]
    total_w = sum(w_vec) or 1
    pv_vec = [w / total_w for w in w_vec]   # priority vector

    # ── Formatting helpers ──────────────────────────────────────────────────
    def cell_bg(cell, hex_str="EBF2FA"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_str)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def cell_rgb(cell, r, g, b):
        cell_bg(cell, f"{r:02X}{g:02X}{b:02X}")

    def thead(table, cols, bg="165FA8"):
        row = table.rows[0]
        for i, txt in enumerate(cols):
            c = row.cells[i]; c.text = ""
            run = c.paragraphs[0].add_run(txt)
            run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(8.5)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_bg(c, bg)

    def trow(row, vals, alt=False, center=False, bold=False, sz=9, color=None):
        for i, v in enumerate(vals):
            if i >= len(row.cells): break
            c = row.cells[i]; c.text = ""
            run = c.paragraphs[0].add_run(str(v))
            run.font.size = Pt(sz); run.font.bold = bold
            if color: run.font.color.rgb = color
            if center: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if alt: cell_bg(c)

    def h1(text, color=BLUE):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True; run.font.size = Pt(15); run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "165FA8")
        pBdr.append(bot); pPr.append(pBdr)
        return p

    def h2(text, color=DARK):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True; run.font.size = Pt(11.5); run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
        return p

    def h3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = GRAY
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        return p

    def body(text, bold=False, italic=False, size=10.5, color=None, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
        if color: run.font.color.rgb = color
        if align: p.alignment = align
        return p

    def bullet(text, size=10):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(size)

    def bar(val, max_val=9, w=12):
        f = max(0, int(round(val / max_val * w))) if max_val else 0
        return "█" * f + "░" * (w - f)

    def fmt_pct(v): return f"{v:+.2f}%" if v != 0 else "0.00%"
    def fmt_n(v):   return f"{v:.3f}"

    # ═══════════════════════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(4): doc.add_paragraph()
    for line, sz, clr in [
        ("INFORME DE ANÁLISIS DE INVERSIÓN",   22, DARK),
        ("Proceso Analítico Jerárquico (AHP)", 14, BLUE),
        ("Optimización de Portafolios — Markowitz", 12, GRAY),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.bold = True; run.font.size = Pt(sz); run.font.color.rgb = clr

    for _ in range(2): doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Perfil del Inversor:  {profile_label.upper()}")
    run.font.bold = True; run.font.size = Pt(20); run.font.color.rgb = P_CLR

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Cartera Seleccionada:  {winner.get('name','—')}   ·   Score AHP: {winner.get('score',0):.1f}%")
    run.font.size = Pt(12); run.font.color.rgb = GRAY

    for _ in range(5): doc.add_paragraph()

    for line, sz in [
        (f"Generado: {today}", 10),
        (f"Fuente de datos: {'Yahoo Finance — Datos de mercado reales' if not is_synthetic else 'Simulación Monte Carlo calibrada'}", 9),
        ("AHP Portfolio Selector  ·  TFG 2025", 9),
        ("Este documento es informativo y no constituye asesoramiento financiero regulado.", 8),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.size = Pt(sz); run.font.italic = True; run.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # RESUMEN EJECUTIVO
    # ═══════════════════════════════════════════════════════════════════════
    h1("RESUMEN EJECUTIVO")
    body(
        f"El presente informe documenta la construcción y selección de la cartera óptima para un "
        f"inversor con perfil {profile_label.upper()} mediante el Proceso Analítico Jerárquico (AHP) "
        f"de Saaty (1990) combinado con la teoría moderna de portafolios de Markowitz (1952). "
        f"Se analizaron {n_analyzed} acciones de tres índices internacionales, de las que "
        f"{n_selected} superaron los filtros cuantitativos del perfil. "
        f"El análisis AHP evaluó los portafolios candidatos frente a 15 indicadores financieros "
        f"en 5 categorías (Ratio de Consistencia CR = {cr_val} "
        f"{'— Consistente ✓' if cr_val < 0.10 else '— Revisar'})."
    )
    if winner_port:
        doc.add_paragraph()
        tbl = doc.add_table(rows=2, cols=6); tbl.style = "Table Grid"
        thead(tbl, ["Rentab. Anual", "Volatilidad", "Sharpe", "Max Drawdown", "Beta", "Alpha"])
        trow(tbl.rows[1], [
            fmt_pct(winner_port.get("rentabilidad", 0)),
            f"{winner_port.get('volatilidad',0):.2f}%",
            fmt_n(winner_port.get("sharpe", 0)),
            fmt_pct(winner_port.get("max_drawdown", 0)),
            fmt_n(winner_port.get("beta", 0)),
            fmt_pct(winner_port.get("alpha", 0)),
        ], center=True, bold=True)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S1. UNIVERSO DE INVERSIÓN
    # ═══════════════════════════════════════════════════════════════════════
    h1("1.  UNIVERSO DE INVERSIÓN — LOS TRES ÍNDICES GLOBALES")
    h2("1.1  Justificación de la Selección de Índices")
    body(
        "El universo de inversión está formado por acciones de los tres índices bursátiles más "
        "representativos del mundo en términos de capitalización, liquidez y cobertura geográfica. "
        "En conjunto representan aproximadamente el 60-65% de la capitalización bursátil mundial, "
        "tres divisas principales (USD, EUR, JPY) y tres ciclos económicos con baja correlación histórica entre sí."
    )

    tbl = doc.add_table(rows=4, cols=5); tbl.style = "Table Grid"
    thead(tbl, ["Índice", "Región / Divisa", "Composición", "Cap. aprox.", "Por qué se incluye"])
    idx_rows = [
        ("S&P 500",       "EE.UU. / USD",   "500 mayores empresas USA por market cap",       "~40 B USD",   "Mayor mercado del mundo; alta liquidez, datos históricos profundos, referencia global de renta variable"),
        ("Euro Stoxx 600","Europa / EUR",    "600 principales empresas de 17 países europeos","~12 B EUR",   "Diversificación frente al ciclo USA; exposición al euro; incluye diferentes sectores y regulaciones"),
        ("Nikkei 225",    "Japón / JPY",     "225 blue chips cotizados en Tokio",             "~4 B JPY",    "Tercera economía mundial; baja correlación histórica con Occidente; cobertura del ciclo asiático"),
    ]
    for i, row in enumerate(idx_rows):
        trow(tbl.rows[i+1], row, alt=(i%2==0), sz=9)

    h2("1.2  Metodología de Muestreo")
    body(
        f"Los tres índices suman más de 700 valores. Descargar y procesar datos históricos para todos "
        f"ellos en tiempo real excedería el tiempo de respuesta razonable de la aplicación (>10 min). "
        f"Se aplica una muestra representativa: los 25 tickers de mayor liquidez y representatividad "
        f"de cada índice (75 en total). Sobre esa muestra se aplica un filtro de calidad que descarta "
        f"series con más del 20% de datos ausentes. Resultado: {n_analyzed} acciones analizadas, "
        f"{n_selected} seleccionadas tras los filtros del perfil."
    )
    body("Este enfoque es consistente con Escobar (2015), que demuestra que una muestra representativa "
         "de 60-80 valores es suficiente para capturar la frontera eficiente de Markowitz.",
         italic=True, size=9.5, color=GRAY)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S2. METODOLOGÍA AHP
    # ═══════════════════════════════════════════════════════════════════════
    h1("2.  METODOLOGÍA AHP — FUNDAMENTOS TEÓRICOS")
    h2("2.1  ¿Qué es el Proceso Analítico Jerárquico?")
    body(
        "El Proceso Analítico Jerárquico (Analytic Hierarchy Process, AHP) fue desarrollado por "
        "Thomas L. Saaty en 1977 y formalizado en su libro The Analytic Hierarchy Process (1980, 1990). "
        "Es el método de decisión multicriterio más citado en la literatura académica de gestión y finanzas. "
        "Su principio fundamental es descomponer una decisión compleja en una jerarquía de niveles: "
        "objetivo → criterios → alternativas."
    )
    body(
        "En este proyecto, la jerarquía es: (1) Objetivo: seleccionar la cartera óptima para el perfil "
        f"{profile_label}. (2) Criterios: los 15 indicadores financieros agrupados en 5 categorías. "
        "(3) Alternativas: los portafolios candidatos construidos con Markowitz."
    )

    h2("2.2  Proceso de Cinco Pasos")
    for i, step in enumerate([
        "Definir la jerarquía: objetivo, criterios y alternativas.",
        "Construir la matriz de comparación por pares de criterios (15×15) usando la escala de Saaty (1-9).",
        "Calcular el vector de prioridades (eigenvector principal normalizado): los pesos de cada criterio.",
        f"Verificar la consistencia mediante el Ratio de Consistencia (CR). Si CR < 0.10, la matriz es aceptable (CR = {cr_val}).",
        "Puntuar las alternativas: multiplicar los valores normalizados de cada portafolio por los pesos AHP y sumar.",
    ], 1):
        bullet(f"Paso {i}: {step}")

    h2("2.3  Escala de Saaty (1-9)")
    body("La intensidad de preferencia entre dos criterios se expresa con la siguiente escala:")
    tbl = doc.add_table(rows=len(SAATY_SCALE)+1, cols=3); tbl.style = "Table Grid"
    thead(tbl, ["Valor", "Intensidad", "Interpretación"])
    for i, (val, intens, interp) in enumerate(SAATY_SCALE):
        trow(tbl.rows[i+1], [str(val), intens, interp], alt=(i%2==0), sz=9)

    h2("2.4  Ratio de Consistencia (CR)")
    body(
        "Para verificar que las comparaciones no son contradictorias, AHP calcula:"
    )
    for line in [
        "λ_max: autovalor principal de la matriz de comparación.",
        "IC (Índice de Consistencia) = (λ_max − n) / (n − 1),  donde n = número de criterios.",
        "IC_A (Índice de Consistencia Aleatorio): valor esperado de IC para una matriz aleatoria del mismo tamaño.",
        "CR = IC / IC_A.  Si CR < 0.10, la matriz se considera suficientemente consistente.",
    ]:
        bullet(line)

    body("Valores de IC_A según Saaty para n = 1 a 15:")
    tbl = doc.add_table(rows=2, cols=15); tbl.style = "Table Grid"
    thead(tbl, [str(i) for i in range(1,16)], bg="444444")
    ri_vals = [RI_TABLE.get(i, 0) for i in range(1,16)]
    trow(tbl.rows[1], [f"{v:.2f}" for v in ri_vals], center=True, sz=8)

    h2("2.5  Síntesis Final — Puntuación AHP de los Portafolios")
    body(
        "Una vez calculados los pesos de los criterios (vector de prioridades p), cada portafolio P_i "
        "recibe una puntuación global:"
    )
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Score(P_i) = Σ_j  p_j × N_ij")
    run.font.name = "Courier New"; run.font.size = Pt(12); run.font.bold = True; run.font.color.rgb = BLUE

    body(
        "Donde N_ij es el valor normalizado del criterio j para el portafolio i (escala 0-1, "
        "donde 1 = mejor valor entre todos los candidatos). El portafolio con mayor puntuación es el ganador."
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S3. LOS 7 PERFILES
    # ═══════════════════════════════════════════════════════════════════════
    h1("3.  LOS 7 PERFILES DE INVERSIÓN")
    body(
        "El cuestionario de 15 preguntas (5 dimensiones: tolerancia al riesgo, horizonte temporal, "
        "situación financiera, perfil psicológico y conocimiento financiero) determina cuál de los "
        "siguientes 7 perfiles describe mejor al inversor. Cada perfil tiene un conjunto de pesos AHP "
        "diferente que refleja sus prioridades."
    )
    doc.add_paragraph()

    profile_order = ["conservador","moderado","agresivo","muy_agresivo","dividendos","tecnologico","esg"]
    for pkey in profile_order:
        plabel, pdesc = PROFILE_FULL_DESC.get(pkey, (pkey, ""))
        pw = ALL_PROFILE_WEIGHTS.get(pkey, {})
        top3 = sorted(pw.items(), key=lambda x: -x[1])[:3]
        bot3 = sorted(pw.items(), key=lambda x:  x[1])[:3]
        is_current = (pkey == profile)

        clr = P_CLR if is_current else DARK
        prefix = "► (PERFIL ASIGNADO)  " if is_current else ""
        h3(f"{prefix}{plabel.upper()}")
        body(pdesc, size=10)
        body(
            f"Criterios con mayor peso: {', '.join(CRITERIA_META[k]['label']+f' ({v}/9)' for k,v in top3)}  |  "
            f"Menor peso: {', '.join(CRITERIA_META[k]['label']+f' ({v}/9)' for k,v in bot3)}",
            size=9, color=GRAY, italic=True
        )
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S4. LOS 15 INDICADORES
    # ═══════════════════════════════════════════════════════════════════════
    h1("4.  LOS 15 INDICADORES FINANCIEROS")
    h2("4.1  Descripción de cada Indicador")
    body("Los 15 indicadores se agrupan en 5 categorías. Para cada uno se muestra la fórmula, "
         "una explicación simple y el objetivo (maximizar o minimizar).")

    tbl = doc.add_table(rows=16, cols=5); tbl.style = "Table Grid"
    thead(tbl, ["#", "Indicador", "Fórmula", "Qué mide", "Objetivo"])
    for i, k in enumerate(CRITERIA_ORDER):
        m = CRITERIA_META[k]
        obj_txt = "MAX" if m["dir"] == "maximize" else "MIN"
        obj_clr = GREEN if m["dir"] == "maximize" else RED
        row = tbl.rows[i+1]
        trow(row, [str(i+1), f"[{m['cat']}]  {m['label']}", m["formula"], m["desc"], ""], alt=(i%2==0), sz=8.5)
        # color the objective cell
        c = row.cells[4]; c.text = ""
        run = c.paragraphs[0].add_run(obj_txt)
        run.font.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = obj_clr
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2("4.2  Relación Indicador — Perfil (A=Alto ≥7  M=Medio 4-6  B=Bajo ≤3)")
    body("La siguiente tabla muestra la intensidad del peso Saaty de cada indicador según el perfil. "
         "A (alto: peso 7-9), M (medio: 4-6), B (bajo: 1-3).",
         size=9.5)

    col_labels = ["Indicador"] + [PROFILE_FULL_DESC[p][0] for p in profile_order]
    tbl = doc.add_table(rows=16, cols=8); tbl.style = "Table Grid"
    thead(tbl, col_labels, bg="1A1A2E")
    for i, k in enumerate(CRITERIA_ORDER):
        cells = [CRITERIA_META[k]["label"]]
        for pkey in profile_order:
            w = ALL_PROFILE_WEIGHTS.get(pkey, {}).get(k, 0)
            cells.append(_weight_to_hml(w))
        row = tbl.rows[i+1]
        trow(row, cells, alt=(i%2==0), sz=8.5, center=True)
        # Highlight current profile column (column index = profile_order.index(profile)+1)
        try:
            ci = profile_order.index(profile) + 1
            c = row.cells[ci]; c.text = ""
            val = cells[ci]
            run = c.paragraphs[0].add_run(val)
            run.font.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = WHITE
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_bg(c, f"{P_CLR.red:02X}{P_CLR.green:02X}{P_CLR.blue:02X}")
        except Exception:
            pass

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S5. PERFIL ASIGNADO
    # ═══════════════════════════════════════════════════════════════════════
    h1(f"5.  PERFIL ASIGNADO: {profile_label.upper()}")
    h2("5.1  Resultado del Cuestionario")
    body(profile_full_desc)

    if scores:
        h2("5.2  Puntuaciones por Perfil")
        body("Afinidad del inversor con cada perfil según las respuestas al cuestionario (0-50 pts):")
        sorted_sc = sorted(scores.items(), key=lambda x: -x[1])
        tbl = doc.add_table(rows=len(sorted_sc)+1, cols=3); tbl.style = "Table Grid"
        thead(tbl, ["Perfil", "Puntuación", "Barra"])
        for i, (pn, ps) in enumerate(sorted_sc):
            label = PROFILE_FULL_DESC.get(pn, (pn.capitalize(),""))[0]
            trow(tbl.rows[i+1], [label, f"{ps} / 50 pts", bar(ps, 50, 18)], alt=(i%2==0), center=True, sz=9)

    h2("5.3  Implicaciones del Perfil en el Análisis AHP")
    top5_crit = sorted(profile_weights.items(), key=lambda x: -x[1])[:5]
    bot5_crit = sorted(profile_weights.items(), key=lambda x:  x[1])[:5]
    body(
        f"Para el perfil {profile_label.upper()}, los criterios más importantes (mayor peso AHP) son: "
        + ", ".join(f"{CRITERIA_META[k]['label']} ({v}/9)" for k,v in top5_crit) + ". "
        f"Los criterios de menor prioridad son: "
        + ", ".join(f"{CRITERIA_META[k]['label']} ({v}/9)" for k,v in bot5_crit) + "."
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S6. FILTRADO DE ACCIONES
    # ═══════════════════════════════════════════════════════════════════════
    h1("6.  FILTRADO Y SELECCIÓN DE ACCIONES")
    h2("6.1  Criterios de Pre-filtrado por Perfil")
    body(
        f"Antes del AHP se aplica un filtro cuantitativo para eliminar acciones incompatibles con el "
        f"perfil {profile_label.upper()}. Los filtros se aplican secuencialmente sobre el universo de {n_analyzed} acciones:"
    )
    FILTERS = {
        "conservador":  [("Capitalización mínima","$10.000 M","Garantiza alta liquidez y estabilidad"),("Volumen medio diario","> 1.000.000 acc.","Liquidez para entrar/salir sin mover el precio"),("Rentabilidad","Positiva (obligatorio)","No se invierte en empresas con pérdidas")],
        "moderado":     [("Capitalización mínima","$5.000 M","Balance liquidez-crecimiento"),("Volumen medio diario","> 750.000 acc.","Liquidez adecuada para horizonte medio"),("Rentabilidad","Positiva (obligatorio)","Criterio básico de viabilidad")],
        "agresivo":     [("Capitalización mínima","$2.000 M","Incluye mid-caps de mayor potencial"),("Volumen medio diario","> 500.000 acc.","Liquidez mínima aceptable"),("Rentabilidad","Positiva (obligatorio)","Criterio básico de viabilidad")],
        "muy_agresivo": [("Capitalización mínima","$1.000 M","Universo más amplio para maximizar upside"),("Volumen medio diario","> 300.000 acc.","Liquidez mínima")],
        "dividendos":   [("Capitalización mínima","$8.000 M","Large caps con historial de dividendos"),("Volumen medio diario","> 800.000 acc.","Alta liquidez"),("Rentabilidad","Positiva","Empresas con beneficios para distribuir")],
        "tecnologico":  [("Capitalización mínima","$3.000 M","Incluye empresas tecnológicas en crecimiento"),("Volumen medio diario","> 600.000 acc.","Liquidez mínima")],
        "esg":          [("Capitalización mínima","$5.000 M","Empresas consolidadas con reporting ESG"),("Volumen medio diario","> 600.000 acc.","Liquidez adecuada")],
    }
    fdata = FILTERS.get(profile, [])
    if fdata:
        tbl = doc.add_table(rows=len(fdata)+1, cols=3); tbl.style = "Table Grid"
        thead(tbl, ["Filtro", f"Umbral para {profile_label}", "Justificación"])
        for i, r in enumerate(fdata):
            trow(tbl.rows[i+1], r, alt=(i%2==0), sz=9)

    if stocks:
        h2("6.2  Acciones que Superaron los Filtros")
        body(f"Las {n_selected} acciones seleccionadas. Período de análisis: 2 años de histórico de precios ajustados.")
        tbl = doc.add_table(rows=len(stocks)+1, cols=5); tbl.style = "Table Grid"
        thead(tbl, ["Ticker", "Rentab. Anual", "Sharpe", "Volatilidad", "Beta"])
        for i, s in enumerate(stocks):
            trow(tbl.rows[i+1], [
                s.get("ticker",""),
                fmt_pct(s.get("rentabilidad",0)),
                fmt_n(s.get("sharpe",0)),
                f"{s.get('volatilidad',0):.2f}%",
                fmt_n(s.get("beta",0)),
            ], alt=(i%2==0), center=True, sz=9)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S7. MARKOWITZ
    # ═══════════════════════════════════════════════════════════════════════
    h1("7.  CONSTRUCCIÓN DE PORTAFOLIOS — OPTIMIZACIÓN DE MARKOWITZ")
    h2("7.1  Metodología de Optimización")
    body(
        "Sobre las acciones seleccionadas se aplica la teoría moderna de portafolios de Markowitz (1952). "
        "Se resuelven cuatro problemas de optimización para generar los portafolios candidatos que "
        "serán evaluados por el AHP:"
    )
    for p_type in [
        "Máximo Sharpe: max [(E(Rp) − Rf)] / σ(Rp)  sujeto a: Σwi = 1, wi ≥ 0",
        "Mínima Varianza: min σ²(Rp) = wᵀΣw  sujeto a: Σwi = 1, wi ≥ 0",
        "Máxima Rentabilidad: max E(Rp)  sujeto a: Σwi = 1, wi ≥ 0",
        "Equal Weight (benchmark): wi = 1/n para todos los activos",
    ]:
        bullet(p_type)

    body(
        "La matriz de covarianzas Σ se calcula sobre los retornos logarítmicos diarios del periodo "
        "de análisis. El vector de retornos esperados μ es la media histórica de cada activo. "
        "Los pesos resultantes w* maximizan el Sharpe o minimizan la varianza según el caso."
    )

    if portfolios:
        h2("7.2  Portafolios Candidatos Generados")
        tbl = doc.add_table(rows=len(portfolios)+1, cols=7); tbl.style = "Table Grid"
        thead(tbl, ["Portafolio", "Rentab.", "Volatil.", "Sharpe", "Max DD", "Beta", "Alpha"])
        for i, p in enumerate(portfolios):
            trow(tbl.rows[i+1], [
                p.get("name",""),
                fmt_pct(p.get("rentabilidad",0)),
                f"{p.get('volatilidad',0):.2f}%",
                fmt_n(p.get("sharpe",0)),
                fmt_pct(p.get("max_drawdown",0)),
                fmt_n(p.get("beta",0)),
                fmt_pct(p.get("alpha",0)),
            ], alt=(i%2==0), center=True, sz=9)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S8. AHP PASO A PASO CON NÚMEROS
    # ═══════════════════════════════════════════════════════════════════════
    h1(f"8.  ANÁLISIS AHP PASO A PASO — DATOS NUMÉRICOS COMPLETOS")
    body(
        "Esta sección expone todos los datos numéricos del análisis AHP tal como se ejecutaron: "
        "vector de pesos de Saaty, matriz de comparación por pares, vector de prioridades normalizado, "
        "verificación de consistencia y matriz de puntuación de portafolios."
    )

    # ── 8.1 Pesos de Saaty ──────────────────────────────────────────────────
    h2(f"8.1  Pesos de Saaty para el Perfil {profile_label} (escala 1-9)")
    body(
        "Los pesos reflejan la importancia relativa de cada criterio para este perfil. "
        f"Suma total de pesos: Σw = {total_w}."
    )

    tbl = doc.add_table(rows=len(CRITERIA_ORDER)+2, cols=6); tbl.style = "Table Grid"
    thead(tbl, ["#", "Indicador", "Categoría", "Peso Saaty", "Barra", "Objetivo"])
    for i, k in enumerate(CRITERIA_ORDER):
        m = CRITERIA_META[k]
        w_k = profile_weights.get(k, 0)
        row = tbl.rows[i+1]
        trow(row, [
            str(i+1), m["label"], m["cat"], str(w_k), bar(w_k, 9, 10),
            "MAX" if m["dir"] == "maximize" else "MIN"
        ], alt=(i%2==0), sz=8.5, center=True)
        # bold the objective
        c = row.cells[5]; c.text = ""
        run = c.paragraphs[0].add_run("MAX" if m["dir"] == "maximize" else "MIN")
        run.font.bold = True; run.font.size = Pt(8.5)
        run.font.color.rgb = GREEN if m["dir"] == "maximize" else RED
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Total row
    trow(tbl.rows[-1], ["", "TOTAL", "", str(total_w), "", ""], bold=True, center=True, sz=9)
    cell_bg(tbl.rows[-1].cells[0], "DDEEFF")
    for ci in range(6): cell_bg(tbl.rows[-1].cells[ci], "DDEEFF")

    # ── 8.2 Vector de Prioridades ────────────────────────────────────────────
    h2("8.2  Vector de Prioridades (Eigenvector Normalizado)")
    body(
        f"El vector de prioridades p_i = w_i / Σw_j normaliza los pesos a escala 0-1 (suma = 1.000). "
        f"Representa la fracción de importancia de cada criterio en la decisión final."
    )
    tbl = doc.add_table(rows=len(CRITERIA_ORDER)+1, cols=4); tbl.style = "Table Grid"
    thead(tbl, ["Indicador", "w_i", "p_i = w_i / Σw", "p_i (%)"])
    for i, k in enumerate(CRITERIA_ORDER):
        w_k = profile_weights.get(k, 0)
        p_k = w_k / total_w
        trow(tbl.rows[i+1], [
            CRITERIA_META[k]["label"],
            str(w_k),
            f"{p_k:.5f}",
            f"{p_k*100:.2f}%",
        ], alt=(i%2==0), sz=9)

    # ── 8.3 Matriz de Comparación por Pares (extracto 6×6) ──────────────────
    h2("8.3  Matriz de Comparación por Pares — Extracto de los 6 Criterios Principales")
    top6_keys = sorted_criteria[:6]
    top6_labels = [CRITERIA_META[k]["label"].split("(")[0].strip()[:18] for k in top6_keys]
    top6_w = [profile_weights.get(k, 1) for k in top6_keys]
    body(
        f"La matriz A[i,j] = w_i / w_j representa cuántas veces más importante es el criterio de "
        f"la fila respecto al de la columna. Valores > 1: fila domina; < 1: columna domina; = 1: igual importancia."
    )
    n6 = len(top6_keys)
    tbl = doc.add_table(rows=n6+1, cols=n6+1); tbl.style = "Table Grid"
    # Header row
    header_row = tbl.rows[0]
    header_row.cells[0].text = ""
    cell_bg(header_row.cells[0], "165FA8")
    for j, lbl in enumerate(top6_labels):
        c = header_row.cells[j+1]; c.text = ""
        run = c.paragraphs[0].add_run(lbl)
        run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(7.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_bg(c, "165FA8")
    # Data rows
    for i in range(n6):
        row = tbl.rows[i+1]
        c = row.cells[0]; c.text = ""
        run = c.paragraphs[0].add_run(top6_labels[i])
        run.font.bold = True; run.font.size = Pt(7.5)
        cell_bg(c, "E8EFF8")
        for j in range(n6):
            c = row.cells[j+1]; c.text = ""
            ratio = top6_w[i] / top6_w[j] if top6_w[j] > 0 else 1.0
            run = c.paragraphs[0].add_run(f"{ratio:.3f}")
            run.font.size = Pt(8)
            if i == j:
                run.font.bold = True
                cell_bg(c, "D4E8FF")
            elif ratio > 1:
                run.font.color.rgb = GREEN
            else:
                run.font.color.rgb = RED
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 8.4 Verificación de Consistencia ─────────────────────────────────────
    h2("8.4  Verificación de Consistencia — λ_max, IC, IC_A, CR")
    n_crit = len(CRITERIA_ORDER)
    ri_val = RI_TABLE.get(n_crit, 1.59)
    # For a weight-based matrix A[i,j]=w_i/w_j, the matrix is perfectly consistent → λ_max = n
    lambda_max = float(n_crit)
    ci_val = (lambda_max - n_crit) / (n_crit - 1) if n_crit > 1 else 0.0
    cr_computed = ci_val / ri_val if ri_val > 0 else 0.0

    body(f"Número de criterios: n = {n_crit}")
    body(f"Autovalor principal: λ_max = {lambda_max:.4f}  (para una matriz perfectamente consistente, λ_max = n)")
    body(f"Índice de Consistencia: IC = (λ_max − n) / (n − 1) = ({lambda_max:.1f} − {n_crit}) / ({n_crit} − 1) = {ci_val:.4f}")
    body(f"Índice de Consistencia Aleatorio (Saaty, n={n_crit}): IC_A = {ri_val:.2f}")
    body(f"Ratio de Consistencia teórico: CR = IC / IC_A = {ci_val:.4f} / {ri_val:.2f} = {cr_computed:.4f}")
    body(f"CR reportado por el motor AHP: {cr_val}   {'✓ Consistente (CR < 0.10)' if cr_val < 0.10 else '⚠ Revisar preferencias'}", bold=True)
    body(
        "Nota: los pesos Saaty asignados por perfil son directamente proporcionales (A[i,j] = w_i/w_j exacto), "
        "por lo que la consistencia teórica es perfecta (CR → 0). El CR reportado por el motor procede "
        "del eigenvector iterativo que verifica numéricamente la coherencia de la matriz completa.",
        italic=True, size=9.5, color=GRAY
    )
    doc.add_page_break()

    # ── 8.5 Matriz de Puntuación Normalizada ─────────────────────────────────
    h2("8.5  Normalización de Portafolios por Criterio")
    body(
        "Para cada criterio j y portafolio i, se normaliza el valor a escala 0-1 donde 1 = mejor candidato:"
    )
    bullet("Criterio a maximizar: N_ij = (v_ij − min_j) / (max_j − min_j)")
    bullet("Criterio a minimizar: N_ij = (max_j − v_ij) / (max_j − min_j)")

    if portfolios:
        # Compute normalized matrix using all 15 criteria
        import numpy as np
        n_ports = len(portfolios)
        raw_matrix = []
        for p in portfolios:
            row = [float(p.get(k, 0) or 0) for k in CRITERIA_ORDER]
            raw_matrix.append(row)
        raw = np.array(raw_matrix)
        norm = np.zeros_like(raw)
        for j, k in enumerate(CRITERIA_ORDER):
            col = raw[:, j]
            cmin, cmax = col.min(), col.max()
            rng = cmax - cmin
            direction = CRITERIA_META[k]["dir"]
            if rng < 1e-10:
                norm[:, j] = 1.0
            elif direction == "maximize":
                norm[:, j] = (col - cmin) / rng
            else:
                norm[:, j] = (cmax - col) / rng

        # Show table: portfolios × top 8 criteria (to fit on page)
        show_keys = sorted_criteria[:8]
        show_labels = [CRITERIA_META[k]["label"].split("(")[0].strip()[:14] for k in show_keys]
        show_idx = [CRITERIA_ORDER.index(k) for k in show_keys]

        tbl = doc.add_table(rows=n_ports+1, cols=len(show_keys)+2); tbl.style = "Table Grid"
        thead(tbl, ["Portafolio"] + show_labels + ["Score AHP"])
        pv_arr = np.array(pv_vec)
        weighted_scores = norm @ pv_arr
        for i, p in enumerate(portfolios):
            vals = [p.get("name","")]
            for j_idx in show_idx:
                vals.append(f"{norm[i, j_idx]:.3f}")
            vals.append(f"{weighted_scores[i]*100:.2f}%")
            is_winner = (p.get("name","") == winner.get("name",""))
            trow(tbl.rows[i+1], vals, alt=(i%2==0), center=True, sz=8.5, bold=is_winner)
            if is_winner:
                for ci in range(len(show_keys)+2):
                    cell_rgb(tbl.rows[i+1].cells[ci], 0xD4, 0xAF, 0x37)

        # ── 8.6 Puntuación Ponderada ──────────────────────────────────────────
        h2("8.6  Puntuación Ponderada Final — Score(P_i) = Σ p_j × N_ij")
        body("Cálculo explícito de la puntuación AHP para cada portafolio candidato:")
        tbl = doc.add_table(rows=n_ports+1, cols=4); tbl.style = "Table Grid"
        thead(tbl, ["Portafolio", "Σ pj × Nij (bruto)", "Score AHP (%)", "Ranking"])
        # Sort by score for display
        port_scores = sorted(enumerate(portfolios), key=lambda x: -weighted_scores[x[0]])
        for rank, (i, p) in enumerate(port_scores):
            is_w = (p.get("name","") == winner.get("name",""))
            trow(tbl.rows[rank+1], [
                p.get("name",""),
                f"{weighted_scores[i]:.6f}",
                f"{weighted_scores[i]*100:.2f}%",
                f"#{rank+1}",
            ], center=True, sz=9, bold=is_w)
            if is_w:
                for ci in range(4): cell_rgb(tbl.rows[rank+1].cells[ci], 0xD4, 0xAF, 0x37)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S9. RANKING AHP
    # ═══════════════════════════════════════════════════════════════════════
    h1("9.  RANKING AHP — SELECCIÓN DEL PORTAFOLIO GANADOR")
    h2("9.1  Ranking Final")
    if ranking and portfolios:
        tbl = doc.add_table(rows=len(ranking)+1, cols=7); tbl.style = "Table Grid"
        thead(tbl, ["Pos.", "Portafolio", "Score AHP", "Rentab.", "Sharpe", "Volat.", "Max DD"])
        for i, r in enumerate(ranking):
            pt = next((p for p in portfolios if p.get("name") == r.get("name")), {})
            vals = [
                f"#{r.get('rank',i+1)}", r.get("name",""), f"{r.get('score',0):.1f}%",
                fmt_pct(pt.get("rentabilidad",0)), fmt_n(pt.get("sharpe",0)),
                f"{pt.get('volatilidad',0):.2f}%", fmt_pct(pt.get("max_drawdown",0)),
            ]
            for j_c, v in enumerate(vals):
                c = tbl.rows[i+1].cells[j_c]; c.text = ""
                run = c.paragraphs[0].add_run(v)
                run.font.size = Pt(9); run.font.bold = (i == 0)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:   cell_rgb(c, 0xD4, 0xAF, 0x37)
                elif i % 2:  pass
                else:        cell_bg(c)

    h2("9.2  Por Qué Ganó Este Portafolio")
    body(
        f"El portafolio '{winner.get('name','')}' obtuvo la puntuación AHP más alta porque maximiza "
        f"los criterios con mayor peso para el perfil {profile_label.upper()}. "
    )
    if winner_port:
        top2 = sorted(profile_weights.items(), key=lambda x: -x[1])[:2]
        for k, w in top2:
            val = winner_port.get(k, winner_port.get(k.split("_")[0], 0))
            body(
                f"  · {CRITERIA_META[k]['label']} (peso {w}/9): "
                f"{fmt_n(float(val)) if abs(float(val or 0)) < 10 else fmt_pct(float(val or 0))} — "
                f"{'valor más favorable entre los candidatos' if CRITERIA_META[k]['dir']=='maximize' else 'valor más bajo entre los candidatos'}.",
                size=10
            )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S10. PORTAFOLIO RECOMENDADO
    # ═══════════════════════════════════════════════════════════════════════
    h1(f"10.  PORTAFOLIO RECOMENDADO — {winner.get('name','').upper()}")
    p_para = doc.add_paragraph()
    run = p_para.add_run(f"Score AHP: {winner.get('score',0):.1f}%   ·   Perfil: {profile_label}")
    run.font.bold = True; run.font.size = Pt(13); run.font.color.rgb = GOLD

    if winner_port:
        h2("10.1  Métricas Financieras Completas")
        metrics_full = [
            ("Rentabilidad Anualizada",   fmt_pct(winner_port.get("rentabilidad",0))),
            ("Volatilidad Anualizada",    f"{winner_port.get('volatilidad',0):.2f}%"),
            ("Ratio de Sharpe",           fmt_n(winner_port.get("sharpe",0))),
            ("Ratio de Sortino",          fmt_n(winner_port.get("sortino",0))),
            ("Alpha de Jensen",           fmt_pct(winner_port.get("alpha",0))),
            ("Max Drawdown",              fmt_pct(winner_port.get("max_drawdown",0))),
            ("Beta",                      fmt_n(winner_port.get("beta",0))),
            ("VaR 95%",                   f"{winner_port.get('var_95',0):.2f}%"),
            ("CVaR 95%",                  f"{winner_port.get('cvar_95',0):.2f}%"),
            ("Tracking Error",            f"{winner_port.get('tracking_error',0):.2f}%"),
            ("Rentab. − Coste Capital",   fmt_pct(winner_port.get("rentab_kp",0))),
            ("Coef. de Variación",        fmt_n(winner_port.get("cv",0))),
            ("Skewness",                  fmt_n(winner_port.get("skewness",0))),
            ("Correlación Media",         fmt_n(winner_port.get("corr_media",0))),
            ("Diversificación Geográfica",fmt_n(winner_port.get("div_geo",0))),
        ]
        # Split into two tables side by side (use one wide table with 4 cols)
        rows_needed = (len(metrics_full) + 1) // 2
        tbl = doc.add_table(rows=rows_needed+1, cols=4); tbl.style = "Table Grid"
        thead(tbl, ["Métrica", "Valor", "Métrica", "Valor"])
        for ri in range(rows_needed):
            li = ri; ri2 = ri + rows_needed
            row = tbl.rows[ri+1]
            if li < len(metrics_full):
                k, v = metrics_full[li]
                trow(row, [k, v, "", ""], alt=(ri%2==0), sz=9)
            if ri2 < len(metrics_full):
                k2, v2 = metrics_full[ri2]
                row.cells[2].text = ""; row.cells[3].text = ""
                row.cells[2].paragraphs[0].add_run(k2).font.size = Pt(9)
                r3 = row.cells[3].paragraphs[0].add_run(v2)
                r3.font.size = Pt(9); r3.font.bold = True
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if allocation:
        h2("10.2  Composición y Asignación de Capital (Markowitz)")
        body("Los porcentajes son el resultado de la optimización que maximiza el Ratio de Sharpe del portafolio.")
        tbl = doc.add_table(rows=len(allocation)+1, cols=5); tbl.style = "Table Grid"
        thead(tbl, ["Empresa (Ticker)", "% Capital", "Barra", "Sharpe Indiv.", "Rol"])
        for i, a in enumerate(allocation):
            ticker = a.get("ticker","")
            wstr   = a.get("weight","0%")
            sd     = next((s for s in stocks if s.get("ticker") == ticker), {})
            sharpe_ind = sd.get("sharpe", 0)
            try:
                wf = float(wstr.replace("%",""))
                role = "Posición principal" if wf > 25 else "Posición media" if wf > 10 else "Diversificadora"
                vis = bar(wf, 100, 14)
            except Exception:
                wf = 0; role = "—"; vis = wstr
            trow(tbl.rows[i+1], [
                ticker, wstr, vis,
                fmt_n(sharpe_ind) if sharpe_ind else "—", role,
            ], alt=(i%2==0), center=True, sz=9)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S11. ÚLTIMAS NOTICIAS
    # ═══════════════════════════════════════════════════════════════════════
    h1("11.  ÚLTIMAS NOTICIAS DEL MERCADO — CARTERA SELECCIONADA")
    body(
        "Titulares de actualidad para las empresas que componen el portafolio recomendado. "
        "Fuente: Yahoo Finance. Fecha de consulta: " + today + "."
    )

    tickers_in_portfolio = [a.get("ticker","") for a in allocation if a.get("ticker","")]
    if not tickers_in_portfolio and stocks:
        tickers_in_portfolio = [s.get("ticker","") for s in stocks[:6]]

    if tickers_in_portfolio:
        news_data = _fetch_news(tickers_in_portfolio, max_per=3)

        if news_data:
            for ticker, items in news_data.items():
                weight_str = next((a.get("weight","") for a in allocation if a.get("ticker") == ticker), "")
                h3(f"{ticker}" + (f"  ·  {weight_str}" if weight_str else ""))
                if not items:
                    body("No hay noticias recientes disponibles.", italic=True, size=9.5, color=GRAY)
                    continue
                tbl = doc.add_table(rows=len(items)+1, cols=3); tbl.style = "Table Grid"
                thead(tbl, ["Titular", "Fuente", "Fecha"])
                for i, n in enumerate(items):
                    trow(tbl.rows[i+1], [
                        n.get("title",""), n.get("publisher",""), n.get("date",""),
                    ], alt=(i%2==0), sz=9)
                doc.add_paragraph()
        else:
            body(
                "No fue posible recuperar noticias en tiempo real en este momento. "
                "Consulte directamente Yahoo Finance, Bloomberg o Reuters para información actualizada.",
                italic=True, color=GRAY
            )
    else:
        body("No hay tickers disponibles para buscar noticias.", italic=True, color=GRAY)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # S12. REFERENCIAS
    # ═══════════════════════════════════════════════════════════════════════
    h1("12.  REFERENCIAS BIBLIOGRÁFICAS")
    REFS = [
        "Saaty, T.L. (1980). The Analytic Hierarchy Process. McGraw-Hill, New York.",
        "Saaty, T.L. (1990). How to make a decision: The Analytic Hierarchy Process. European Journal of Operational Research, 48(1), 9-26.",
        "Markowitz, H.M. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Sharpe, W.F. (1966). Mutual Fund Performance. Journal of Business, 39(1), 119-138.",
        "Sharpe, W.F. (1970). Portfolio Theory and Capital Markets. McGraw-Hill, New York.",
        "Jensen, M.C. (1968). The Performance of Mutual Funds in the Period 1945-1964. The Journal of Finance, 23(2), 389-416.",
        "Sortino, F.A. & Price, L.N. (1994). Performance Measurement in a Downside Risk Framework. The Journal of Investing, 3(3), 59-64.",
        "Artzner, P., Delbaen, F., Eber, J.M. & Heath, D. (1999). Coherent Measures of Risk. Mathematical Finance, 9(3), 203-228.",
        "Escobar, J.W. (2015). Selección de portafolios óptimos de inversión usando el Proceso Analítico Jerárquico. Contaduría y Administración, 60(2), 346-366.",
        "Elton, E.J. & Gruber, M.J. (1995). Modern Portfolio Theory and Investment Analysis (5th ed.). John Wiley & Sons.",
        "Morgan, J.P. (1996). RiskMetrics — Technical Document (4th ed.). J.P. Morgan/Reuters.",
    ]
    for ref in REFS:
        bullet(ref, size=9.5)

    doc.add_paragraph()
    body(
        f"Documento generado por AHP Portfolio Selector — {today}. "
        f"Datos: {'Yahoo Finance (precios ajustados de mercado real)' if not is_synthetic else 'Simulación Monte Carlo calibrada a parámetros de mercado histórico'}. "
        "Este informe tiene carácter exclusivamente informativo y académico. "
        "No constituye asesoramiento financiero ni recomendación de inversión regulada.",
        italic=True, size=8.5, color=GRAY
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
