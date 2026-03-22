"""
Aplicación web — AHP Portfolio Selector (versión PythonAnywhere)
=================================================================
Versión adaptada para despliegue en pythonanywhere.com
"""

import sys
import os
import json
import numpy as np
import pandas as pd
from datetime import datetime

# Path setup para PythonAnywhere
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.join(PROJECT_DIR, "..")
if PARENT_DIR not in sys.path:
    sys.path.insert(0, PARENT_DIR)
if PROJECT_DIR not in sys.path:
    sys.path.insert(0, PROJECT_DIR)

from flask import Flask, render_template, request, jsonify, send_file
from questionnaire import QUESTIONS, PROFILES_ORDER, PROFILE_DESCRIPTIONS
from profiles import get_profile_config, CRITERIA, CRITERIA_ORDER, PROFILE_WEIGHTS
from stock_analysis import compute_returns, analyze_stocks
from stock_filter import run_filtering
from portfolio_builder import build_portfolios
from ahp_engine_v2 import AHPEngine

app = Flask(__name__)
app.config["SECRET_KEY"] = "ahp-portfolio-tfg-2025"


def generate_synthetic_data():
    np.random.seed(int(datetime.now().timestamp()) % 10000)
    n_days = 504
    dates = pd.bdate_range(start="2023-01-02", periods=n_days)

    configs = {
        "sp500": {
            "AAPL": (0.12, 0.25), "MSFT": (0.18, 0.28), "NVDA": (0.25, 0.40),
            "JNJ": (0.05, 0.14), "PG": (0.04, 0.12), "JPM": (0.10, 0.22),
            "XOM": (0.08, 0.24), "KO": (0.03, 0.11),
        },
        "eurostoxx": {
            "SAP.DE": (0.14, 0.26), "SIE.DE": (0.08, 0.20), "SAN.MC": (0.06, 0.24),
            "MC.PA": (0.11, 0.22), "TTE.PA": (0.07, 0.26), "ASML.AS": (0.20, 0.32),
        },
        "nikkei": {
            "7203.T": (0.09, 0.22), "6758.T": (0.15, 0.30), "9984.T": (0.20, 0.35),
            "4502.T": (0.04, 0.16), "8306.T": (0.07, 0.20),
        },
    }

    all_prices, all_benchmarks = {}, {}
    for idx, stocks in configs.items():
        d = {}
        for t, (mu, sig) in stocks.items():
            r = np.random.normal(mu/252, sig/np.sqrt(252), n_days)
            d[t] = 100 * np.exp(np.cumsum(r))
        all_prices[idx] = pd.DataFrame(d, index=dates)
        bmu = np.mean([s[0] for s in stocks.values()])
        bsig = np.mean([s[1] for s in stocks.values()]) * 0.7
        br = np.random.normal(bmu/252, bsig/np.sqrt(252), n_days)
        all_benchmarks[idx] = pd.Series(1000*np.exp(np.cumsum(br)), index=dates)

    return {
        "prices": all_prices, "benchmarks": all_benchmarks,
        "risk_free": {"sp500": 0.04, "eurostoxx": 0.025, "nikkei": 0.01},
        "stock_info": None,
        "metadata": {"indices": list(configs.keys()), "synthetic": True},
    }


def load_live_data():
    try:
        from data_loader import load_market_data
        data = load_market_data(period_years=2, max_per_index=15, progress=False)
        data["metadata"]["synthetic"] = False
        return data
    except Exception as e:
        print(f"Live data failed ({e}), using synthetic")
        return generate_synthetic_data()


@app.route("/")
def index():
    return render_template("index.html", questions=QUESTIONS)


@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        data = request.get_json()
        answers = data.get("answers", {})

        scores = {p: 0 for p in PROFILES_ORDER}
        for q in QUESTIONS:
            resp = answers.get(q["id"], "b")
            for letra, texto, puntos in q["opciones"]:
                if letra == resp:
                    for perfil, pts in puntos.items():
                        scores[perfil] += pts

        profile = max(scores, key=scores.get)
        config = get_profile_config(profile)

        use_live = data.get("use_live", True)
        market_data = load_live_data() if use_live else generate_synthetic_data()
        is_synthetic = market_data["metadata"].get("synthetic", True)

        all_analyses, all_returns = [], {}
        for idx in market_data["metadata"]["indices"]:
            prices = market_data["prices"][idx]
            benchmark = market_data["benchmarks"][idx]
            rf = market_data["risk_free"][idx]
            analysis = analyze_stocks(prices, benchmark, rf, idx, progress=False)
            analysis["region"] = idx
            all_analyses.append(analysis)
            all_returns[idx] = compute_returns(prices)

        consolidated = pd.concat(all_analyses)

        avg_rf = np.mean(list(market_data["risk_free"].values()))
        market_vars = []
        for idx in market_data["metadata"]["indices"]:
            b = market_data["benchmarks"][idx]
            br = np.log(b/b.shift(1)).dropna()
            market_vars.append(br.var()*252)

        selected = run_filtering(
            consolidated, config["filters"], avg_rf,
            np.mean(market_vars), stock_info=None,
            profile_name=profile, progress=False,
        )

        all_ret = pd.concat(all_returns.values(), axis=1)
        first_idx = market_data["metadata"]["indices"][0]
        bench = market_data["benchmarks"][first_idx]
        bench_ret = np.log(bench/bench.shift(1)).dropna()

        best = build_portfolios(selected, all_ret, bench_ret, avg_rf, progress=False)

        engine = AHPEngine(profile=profile)
        ranking = engine.run(best, progress=False)

        weights = PROFILE_WEIGHTS[profile]
        sorted_criteria = sorted(weights.items(), key=lambda x: -x[1])
        top_criteria = [
            {"name": CRITERIA[c]["label"], "weight": w, "direction": CRITERIA[c]["dir"]}
            for c, w in sorted_criteria[:6]
        ]

        ranking_list = [
            {"name": row["portafolio"], "score": round(row["score_pct"], 1), "rank": int(row["ranking"])}
            for _, row in ranking.iterrows()
        ]

        portfolios_detail = []
        for pname, prow in best.iterrows():
            portfolios_detail.append({
                "name": pname,
                "rentabilidad": round(float(prow.get("rentabilidad", 0)) * 100, 2),
                "volatilidad": round(float(prow.get("volatilidad", 0)) * 100, 2),
                "sharpe": round(float(prow.get("sharpe", 0)), 2),
                "max_drawdown": round(float(prow.get("max_drawdown", 0)) * 100, 2),
                "beta": round(float(prow.get("beta", 0)), 2),
                "alpha": round(float(prow.get("alpha", 0)) * 100, 2),
                "tickers": str(prow.get("tickers", "")),
                "pesos": str(prow.get("pesos", "")),
            })

        stocks_list = [
            {"ticker": t, "rentabilidad": round(float(s["rentabilidad"])*100, 2),
             "sharpe": round(float(s["sharpe"]), 2), "volatilidad": round(float(s["volatilidad"])*100, 2),
             "beta": round(float(s["beta"]), 2)}
            for t, s in selected.head(8).iterrows()
        ]

        winner = ranking_list[0]
        winner_name = ranking.iloc[0]["portafolio"]
        winner_data = best.loc[winner_name]
        allocation = []
        if "pesos" in winner_data and isinstance(winner_data["pesos"], str):
            for pair in winner_data["pesos"].split(", "):
                parts = pair.split(":")
                if len(parts) == 2:
                    allocation.append({"ticker": parts[0].strip(), "weight": parts[1].strip()})

        app.config["LAST_RESULT"] = {
            "profile": profile, "profile_description": PROFILE_DESCRIPTIONS[profile],
            "scores": {k: v for k, v in scores.items() if v > 0},
            "top_criteria": top_criteria, "ranking": ranking_list,
            "portfolios": portfolios_detail, "stocks": stocks_list,
            "winner": winner, "allocation": allocation,
            "n_stocks_analyzed": len(consolidated), "n_stocks_selected": len(selected),
            "consistency_ratio": round(engine.criteria_cr, 4),
            "is_synthetic": is_synthetic, "answers": answers,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M"),
        }

        return jsonify({
            "success": True, "profile": profile,
            "profile_description": PROFILE_DESCRIPTIONS[profile],
            "scores": {k: v for k, v in scores.items() if v > 0},
            "top_criteria": top_criteria, "ranking": ranking_list,
            "portfolios": portfolios_detail, "stocks": stocks_list,
            "winner": winner, "allocation": allocation,
            "n_stocks_analyzed": len(consolidated),
            "n_stocks_selected": len(selected),
            "consistency_ratio": round(engine.criteria_cr, 4),
            "is_synthetic": is_synthetic,
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})


@app.route("/download-report")
def download_report():
    data = app.config.get("LAST_RESULT")
    if not data:
        return "No hay resultados. Completa el cuestionario primero.", 400
    try:
        report_path = generate_word_report(data)
        return send_file(
            report_path, as_attachment=True,
            download_name=f"Informe_AHP_{data['profile']}_{datetime.now().strftime('%Y%m%d')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return f"Error generando informe: {e}", 500


def generate_word_report(data):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import tempfile

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME DE SELECCIÓN DE PORTAFOLIO")
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proceso Analítico Jerárquico (AHP)")
    run.font.size = Pt(16); run.font.color.rgb = RGBColor(46, 80, 144)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generado: {data['timestamp']}").font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    profile_name = data["profile"].replace("_", " ").upper()

    h = doc.add_heading("1. Perfil del inversor", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Perfil asignado: {profile_name}")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(46, 80, 144)
    doc.add_paragraph(data["profile_description"])
    doc.add_paragraph("Este perfil ha sido determinado mediante un cuestionario de 12 preguntas que evalúa el horizonte temporal, la tolerancia al riesgo, los objetivos de inversión, la experiencia del inversor y sus preferencias sectoriales.")

    scores_sorted = sorted(data["scores"].items(), key=lambda x: -x[1])
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Perfil"
    table.rows[0].cells[1].text = "Puntuación"
    for pn, ps in scores_sorted:
        row = table.add_row()
        row.cells[0].text = pn.replace("_", " ").capitalize()
        row.cells[1].text = str(ps)

    doc.add_page_break()
    doc.add_heading("2. Índices utilizados y universo de acciones", level=1)
    doc.add_paragraph("El análisis se ha realizado sobre acciones de tres índices bursátiles internacionales:")
    for idx_n, reg, desc in [("S&P 500","EE.UU.","500 mayores empresas estadounidenses"),("Eurostoxx 600","Europa","600 principales empresas europeas"),("Nikkei 225","Japón","225 empresas representativas de Tokio")]:
        doc.add_paragraph(f"{idx_n} ({reg}): {desc}", style="List Bullet")
    is_real = not data.get("is_synthetic", True)
    doc.add_paragraph(f"Total acciones analizadas: {data['n_stocks_analyzed']}. Seleccionadas: {data['n_stocks_selected']}.")
    doc.add_paragraph(f"Fuente: {'Yahoo Finance (datos reales)' if is_real else 'Datos sintéticos'}.")

    doc.add_heading("3. Acciones seleccionadas", level=1)
    doc.add_paragraph(f"Acciones seleccionadas tras el filtrado para el perfil {profile_name}:")
    if data.get("stocks"):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h_text in enumerate(["Ticker", "Rentabilidad", "Sharpe", "Volatilidad"]):
            table.rows[0].cells[i].text = h_text
        for s in data["stocks"]:
            row = table.add_row()
            row.cells[0].text = s["ticker"]
            row.cells[1].text = f"{s['rentabilidad']}%"
            row.cells[2].text = str(s["sharpe"])
            row.cells[3].text = f"{s['volatilidad']}%"

    doc.add_heading("4. Criterios AHP aplicados", level=1)
    doc.add_paragraph(f"Para el perfil {profile_name}, los criterios con mayor peso son:")
    if data.get("top_criteria"):
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Criterio"
        table.rows[0].cells[1].text = "Peso (1-9)"
        table.rows[0].cells[2].text = "Dirección"
        for c in data["top_criteria"]:
            row = table.add_row()
            row.cells[0].text = c["name"]
            row.cells[1].text = str(c["weight"])
            row.cells[2].text = "Maximizar" if c["direction"] == "maximize" else "Minimizar"
    doc.add_paragraph(f"Ratio de consistencia: {data['consistency_ratio']} (< 0.10 = consistente).")

    doc.add_page_break()
    doc.add_heading("5. Portafolio recomendado", level=1)
    winner = data["winner"]
    p = doc.add_paragraph()
    run = p.add_run(f"Portafolio ganador: {winner['name']} — Puntuación AHP: {winner['score']}%")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(84, 130, 53)

    if data.get("allocation"):
        doc.add_heading("5.1. Asignación de inversión", level=2)
        doc.add_paragraph("Porcentaje del capital a invertir en cada acción (optimizado por Markowitz):")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Acción"
        table.rows[0].cells[1].text = "% del capital"
        for a in data["allocation"]:
            row = table.add_row()
            row.cells[0].text = a["ticker"]
            row.cells[1].text = a["weight"]

    doc.add_heading("5.2. Ranking completo", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Ranking"
    table.rows[0].cells[1].text = "Portafolio"
    table.rows[0].cells[2].text = "Puntuación"
    for r in data["ranking"]:
        row = table.add_row()
        row.cells[0].text = f"#{r['rank']}"
        row.cells[1].text = r["name"]
        row.cells[2].text = f"{r['score']}%"

    if data.get("portfolios"):
        doc.add_heading("5.3. Métricas de portafolios", level=2)
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        for i, h_text in enumerate(["Portaf.", "Rentab.", "Volat.", "Sharpe", "MaxDD", "Beta", "Alpha"]):
            table.rows[0].cells[i].text = h_text
        for pd_item in data["portfolios"]:
            row = table.add_row()
            row.cells[0].text = pd_item["name"]
            row.cells[1].text = f"{pd_item['rentabilidad']}%"
            row.cells[2].text = f"{pd_item['volatilidad']}%"
            row.cells[3].text = str(pd_item["sharpe"])
            row.cells[4].text = f"{pd_item['max_drawdown']}%"
            row.cells[5].text = str(pd_item["beta"])
            row.cells[6].text = f"{pd_item['alpha']}%"

    doc.add_page_break()
    doc.add_heading("6. Nota metodológica", level=1)
    doc.add_paragraph("Este análisis utiliza la técnica AHP (Saaty, 1990) adaptada a portafolios de inversión según Escobar (2015), con 15 criterios en 5 categorías, optimización de Markowitz (1952) y filtrado por perfil con la teoría de la razón beta de Elton y Gruber (1978).")
    doc.add_paragraph("Referencias:", style="Heading 2")
    for ref in ["Saaty, T.L. (1990). How to make a decision: The AHP. EJOR, 48(1), 9-26.",
                "Escobar, J.W. (2015). Metodología AHP para inversión en portafolio. Contaduría y Adm., 60, 346-366.",
                "Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77-91.",
                "Sharpe, W.F. (1966). Mutual fund performance. Journal of Business, 39(1), 119-138.",
                "Sortino, F.A. y Price, L.N. (1994). Performance measurement in a downside risk framework.",
                "Artzner, P. et al. (1999). Coherent measures of risk. Mathematical Finance, 9(3), 203-228."]:
        doc.add_paragraph(ref, style="List Bullet")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
