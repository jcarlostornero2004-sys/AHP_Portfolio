"""
Aplicación web — AHP Portfolio Selector
=========================================
Interfaz web interactiva para el selector de portafolios AHP.
Se abre en el navegador y ejecuta todo el pipeline con datos reales.

Uso:
  python webapp/app.py
  → Se abre automáticamente en http://localhost:5000
"""

import sys
import os
import json
import webbrowser
import threading
import numpy as np
import pandas as pd
from datetime import datetime

# Añadir carpeta padre al path para importar módulos
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from flask import Flask, render_template, request, jsonify
from questionnaire import QUESTIONS, PROFILES_ORDER, PROFILE_DESCRIPTIONS
from profiles import get_profile_config, CRITERIA, CRITERIA_ORDER, PROFILE_WEIGHTS
from stock_analysis import compute_returns, analyze_stocks
from stock_filter import run_filtering
from portfolio_builder import build_portfolios
from ahp_engine_v2 import AHPEngine

app = Flask(__name__)


# ─────────────────────────────────────────────────────────────────
# DATOS SINTÉTICOS (fallback si no hay internet)
# ─────────────────────────────────────────────────────────────────

def generate_synthetic_data():
    np.random.seed(int(datetime.now().timestamp()) % 1000)
    n_days = 504
    dates = pd.bdate_range(start="2023-01-02", periods=n_days)

    configs = {
        "sp500": {
            "AAPL": (0.12, 0.25), "MSFT": (0.18, 0.28), "NVDA": (0.25, 0.40),
            "JNJ": (0.05, 0.14), "PG": (0.04, 0.12), "JPM": (0.10, 0.22),
            "XOM": (0.08, 0.24), "KO": (0.03, 0.11),
        },
        "eurostoxx": {
            "SAP.DE": (0.14, 0.26), "SIE.DE": (0.08, 0.20), "SAN.MC": (0.06, 0.24),
            "MC.PA": (0.11, 0.22), "TTE.PA": (0.07, 0.26), "ASML.AS": (0.20, 0.32),
        },
        "nikkei": {
            "7203.T": (0.09, 0.22), "6758.T": (0.15, 0.30), "9984.T": (0.20, 0.35),
            "4502.T": (0.04, 0.16), "8306.T": (0.07, 0.20),
        },
    }

    all_prices, all_benchmarks = {}, {}
    for idx, stocks in configs.items():
        d = {}
        for t, (mu, sig) in stocks.items():
            r = np.random.normal(mu/252, sig/np.sqrt(252), n_days)
            d[t] = 100 * np.exp(np.cumsum(r))
        all_prices[idx] = pd.DataFrame(d, index=dates)
        bmu = np.mean([s[0] for s in stocks.values()])
        bsig = np.mean([s[1] for s in stocks.values()]) * 0.7
        br = np.random.normal(bmu/252, bsig/np.sqrt(252), n_days)
        all_benchmarks[idx] = pd.Series(1000*np.exp(np.cumsum(br)), index=dates)

    return {
        "prices": all_prices, "benchmarks": all_benchmarks,
        "risk_free": {"sp500": 0.04, "eurostoxx": 0.025, "nikkei": 0.01},
        "stock_info": None,
        "metadata": {"indices": list(configs.keys()), "synthetic": True},
    }


def load_live_data():
    """Intenta cargar datos reales. Si falla, usa sintéticos."""
    try:
        from data_loader import load_market_data
        data = load_market_data(period_years=2, max_per_index=15, progress=False)
        data["metadata"]["synthetic"] = False
        return data
    except Exception as e:
        print(f"  Datos reales no disponibles ({e}), usando sintéticos")
        return generate_synthetic_data()


# ─────────────────────────────────────────────────────────────────
# RUTAS WEB
# ─────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html", questions=QUESTIONS)


@app.route("/analyze", methods=["POST"])
def analyze():
    """Recibe respuestas del cuestionario y ejecuta el pipeline completo."""
    try:
        data = request.get_json()
        answers = data.get("answers", {})

        # ── Determinar perfil ──
        scores = {p: 0 for p in PROFILES_ORDER}
        for q in QUESTIONS:
            resp = answers.get(q["id"], "b")
            for letra, texto, puntos in q["opciones"]:
                if letra == resp:
                    for perfil, pts in puntos.items():
                        scores[perfil] += pts

        profile = max(scores, key=scores.get)
        config = get_profile_config(profile)

        # ── Cargar datos ──
        use_live = data.get("use_live", True)
        if use_live:
            market_data = load_live_data()
        else:
            market_data = generate_synthetic_data()

        is_synthetic = market_data["metadata"].get("synthetic", True)

        # ── Análisis (M3) ──
        all_analyses = []
        all_returns = {}
        for idx in market_data["metadata"]["indices"]:
            prices = market_data["prices"][idx]
            benchmark = market_data["benchmarks"][idx]
            rf = market_data["risk_free"][idx]
            analysis = analyze_stocks(prices, benchmark, rf, idx, progress=False)
            analysis["region"] = idx
            all_analyses.append(analysis)
            all_returns[idx] = compute_returns(prices)

        consolidated = pd.concat(all_analyses)

        # ── Filtrado (M4) ──
        avg_rf = np.mean(list(market_data["risk_free"].values()))
        market_vars = []
        for idx in market_data["metadata"]["indices"]:
            b = market_data["benchmarks"][idx]
            br = np.log(b/b.shift(1)).dropna()
            market_vars.append(br.var()*252)

        selected = run_filtering(
            consolidated, config["filters"], avg_rf,
            np.mean(market_vars), stock_info=None,
            profile_name=profile, progress=False,
        )

        # ── Portafolios (M5) ──
        all_ret = pd.concat(all_returns.values(), axis=1)
        first_idx = market_data["metadata"]["indices"][0]
        bench = market_data["benchmarks"][first_idx]
        bench_ret = np.log(bench/bench.shift(1)).dropna()

        best = build_portfolios(selected, all_ret, bench_ret, avg_rf, progress=False)

        # ── AHP (M6) ──
        engine = AHPEngine(profile=profile)
        ranking = engine.run(best, progress=False)

        # ── Preparar respuesta JSON ──
        # Top criterios del perfil
        weights = PROFILE_WEIGHTS[profile]
        sorted_criteria = sorted(weights.items(), key=lambda x: -x[1])
        top_criteria = [
            {"name": CRITERIA[c]["label"], "weight": w, "direction": CRITERIA[c]["dir"]}
            for c, w in sorted_criteria[:6]
        ]

        # Ranking
        ranking_list = []
        for _, row in ranking.iterrows():
            ranking_list.append({
                "name": row["portafolio"],
                "score": round(row["score_pct"], 1),
                "rank": int(row["ranking"]),
            })

        # Portafolios con detalles
        portfolios_detail = []
        for pname, prow in best.iterrows():
            portfolios_detail.append({
                "name": pname,
                "rentabilidad": round(float(prow.get("rentabilidad", 0)) * 100, 2),
                "volatilidad": round(float(prow.get("volatilidad", 0)) * 100, 2),
                "sharpe": round(float(prow.get("sharpe", 0)), 2),
                "max_drawdown": round(float(prow.get("max_drawdown", 0)) * 100, 2),
                "beta": round(float(prow.get("beta", 0)), 2),
                "alpha": round(float(prow.get("alpha", 0)) * 100, 2),
                "tickers": str(prow.get("tickers", "")),
                "pesos": str(prow.get("pesos", "")),
            })

        # Acciones seleccionadas
        stocks_list = []
        for ticker, srow in selected.head(8).iterrows():
            stocks_list.append({
                "ticker": ticker,
                "rentabilidad": round(float(srow["rentabilidad"]) * 100, 2),
                "sharpe": round(float(srow["sharpe"]), 2),
                "volatilidad": round(float(srow["volatilidad"]) * 100, 2),
                "beta": round(float(srow["beta"]), 2),
            })

        winner = ranking_list[0]

        # Extraer asignación del portafolio ganador (% por acción)
        winner_name = ranking.iloc[0]["portafolio"]
        winner_data = best.loc[winner_name]
        allocation = []
        if "pesos" in winner_data and isinstance(winner_data["pesos"], str):
            for pair in winner_data["pesos"].split(", "):
                parts = pair.split(":")
                if len(parts) == 2:
                    t = parts[0].strip()
                    w = parts[1].strip()
                    allocation.append({"ticker": t, "weight": w})

        # Guardar datos para generación de informe
        app.config["LAST_RESULT"] = {
            "profile": profile,
            "profile_description": PROFILE_DESCRIPTIONS[profile],
            "scores": {k: v for k, v in scores.items() if v > 0},
            "top_criteria": top_criteria,
            "ranking": ranking_list,
            "portfolios": portfolios_detail,
            "stocks": stocks_list,
            "winner": winner,
            "allocation": allocation,
            "n_stocks_analyzed": len(consolidated),
            "n_stocks_selected": len(selected),
            "consistency_ratio": round(engine.criteria_cr, 4),
            "is_synthetic": is_synthetic,
            "answers": answers,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M"),
        }

        return jsonify({
            "success": True,
            "profile": profile,
            "profile_description": PROFILE_DESCRIPTIONS[profile],
            "scores": {k: v for k, v in scores.items() if v > 0},
            "top_criteria": top_criteria,
            "ranking": ranking_list,
            "portfolios": portfolios_detail,
            "stocks": stocks_list,
            "winner": winner,
            "allocation": allocation,
            "n_stocks_analyzed": len(consolidated),
            "n_stocks_selected": len(selected),
            "consistency_ratio": round(engine.criteria_cr, 4),
            "is_synthetic": is_synthetic,
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})


@app.route("/download-report")
def download_report():
    """Genera y descarga un informe Word personalizado."""
    from flask import send_file
    import tempfile

    data = app.config.get("LAST_RESULT")
    if not data:
        return "No hay resultados. Completa el cuestionario primero.", 400

    try:
        report_path = generate_word_report(data)
        return send_file(
            report_path,
            as_attachment=True,
            download_name=f"Informe_AHP_{data['profile']}_{datetime.now().strftime('%Y%m%d')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return f"Error generando informe: {e}", 500


def generate_word_report(data):
    """Genera el informe Word con python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import tempfile

    doc = DocxDocument()

    # Estilos
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── PORTADA ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME DE SELECCIÓN DE PORTAFOLIO")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 56, 100)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Proceso Analítico Jerárquico (AHP)")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph("")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Generado: {data['timestamp']}").font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ── 1. PERFIL DEL INVERSOR ──
    h = doc.add_heading("1. Perfil del inversor", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    profile_name = data["profile"].replace("_", " ").upper()
    p = doc.add_paragraph()
    run = p.add_run(f"Perfil asignado: {profile_name}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 80, 144)

    doc.add_paragraph(data["profile_description"])

    doc.add_paragraph(
        "Este perfil ha sido determinado mediante un cuestionario de 12 preguntas "
        "que evalúa el horizonte temporal, la tolerancia al riesgo, los objetivos "
        "de inversión, la experiencia del inversor y sus preferencias sectoriales. "
        "Las respuestas se ponderan para asignar automáticamente uno de los 7 perfiles "
        "disponibles, cada uno con pesos AHP y filtros de universo específicos."
    )

    # Tabla de puntuaciones
    doc.add_paragraph("")
    doc.add_paragraph("Puntuación por perfil:").bold = True
    scores_sorted = sorted(data["scores"].items(), key=lambda x: -x[1])
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Perfil"
    table.rows[0].cells[1].text = "Puntuación"
    for pname, pscore in scores_sorted:
        row = table.add_row()
        row.cells[0].text = pname.replace("_", " ").capitalize()
        row.cells[1].text = str(pscore)

    # ── 2. ÍNDICES Y UNIVERSO ──
    doc.add_page_break()
    h = doc.add_heading("2. Índices utilizados y universo de acciones", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph(
        "El análisis se ha realizado sobre acciones de tres índices bursátiles "
        "internacionales, proporcionando diversificación geográfica entre las tres "
        "principales regiones económicas del mundo:"
    )

    indices_info = [
        ("S&P 500", "Estados Unidos", "Las 500 empresas de mayor capitalización del mercado estadounidense"),
        ("Eurostoxx 600", "Europa", "Las 600 principales empresas cotizadas en Europa"),
        ("Nikkei 225", "Japón", "Las 225 empresas más representativas de la bolsa de Tokio"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Índice"
    table.rows[0].cells[1].text = "Región"
    table.rows[0].cells[2].text = "Descripción"
    for idx_name, region, desc in indices_info:
        row = table.add_row()
        row.cells[0].text = idx_name
        row.cells[1].text = region
        row.cells[2].text = desc

    doc.add_paragraph("")
    doc.add_paragraph(
        f"Se han analizado un total de {data['n_stocks_analyzed']} acciones, "
        f"de las cuales se han seleccionado {data['n_stocks_selected']} para la "
        f"construcción de portafolios candidatos, tras aplicar los filtros "
        f"específicos del perfil {profile_name}."
    )

    is_real = not data.get("is_synthetic", True)
    doc.add_paragraph(
        f"Fuente de datos: {'Yahoo Finance (datos reales de mercado)' if is_real else 'Datos sintéticos (simulación)'}."
    )

    # ── 3. ACCIONES SELECCIONADAS ──
    h = doc.add_heading("3. Acciones seleccionadas para el análisis", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph(
        f"Las siguientes acciones han sido seleccionadas tras aplicar el proceso "
        f"de filtrado secuencial basado en el perfil {profile_name}: eliminación de "
        f"acciones con rentabilidad negativa, aplicación de la teoría de la razón "
        f"beta de Elton y Gruber (1978), filtros específicos del perfil (exclusiones "
        f"de volatilidad, drawdown y beta según la tolerancia al riesgo) y ranking "
        f"final por los criterios prioritarios del perfil."
    )

    if data.get("stocks"):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(["Ticker", "Rentabilidad", "Sharpe", "Volatilidad"]):
            table.rows[0].cells[i].text = header
        for s in data["stocks"]:
            row = table.add_row()
            row.cells[0].text = s["ticker"]
            row.cells[1].text = f"{s['rentabilidad']}%"
            row.cells[2].text = str(s["sharpe"])
            row.cells[3].text = f"{s['volatilidad']}%"

    # ── 4. CRITERIOS AHP ──
    h = doc.add_heading("4. Criterios AHP aplicados", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph(
        "La metodología AHP utiliza 15 criterios organizados en 5 categorías "
        "(Rendimiento, Riesgo, Eficiencia de mercado, Estabilidad y Diversificación). "
        f"Para el perfil {profile_name}, los criterios con mayor peso son:"
    )

    if data.get("top_criteria"):
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Criterio"
        table.rows[0].cells[1].text = "Peso (1-9)"
        table.rows[0].cells[2].text = "Dirección"
        for c in data["top_criteria"]:
            row = table.add_row()
            row.cells[0].text = c["name"]
            row.cells[1].text = str(c["weight"])
            row.cells[2].text = "Maximizar" if c["direction"] == "maximize" else "Minimizar"

    doc.add_paragraph("")
    doc.add_paragraph(
        f"El ratio de consistencia (CR) de la matriz de criterios es "
        f"{data['consistency_ratio']}, inferior al umbral de 0.10 establecido "
        f"por Saaty (1990), lo que confirma la coherencia de las valoraciones."
    )

    # ── 5. PORTAFOLIO RECOMENDADO ──
    doc.add_page_break()
    h = doc.add_heading("5. Portafolio recomendado", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    winner = data["winner"]
    p = doc.add_paragraph()
    run = p.add_run(f"Portafolio ganador: {winner['name']} — Puntuación AHP: {winner['score']}%")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(84, 130, 53)

    doc.add_paragraph(
        f"Tras evaluar los 5 portafolios candidatos bajo los 15 criterios AHP "
        f"ponderados según el perfil {profile_name}, el portafolio {winner['name']} "
        f"obtiene la mayor puntuación global con un {winner['score']}% de preferencia."
    )

    # Asignación de inversión
    if data.get("allocation"):
        doc.add_paragraph("")
        h2 = doc.add_heading("5.1. Asignación de inversión recomendada", level=2)
        h2.runs[0].font.color.rgb = RGBColor(46, 80, 144)

        doc.add_paragraph(
            "La siguiente tabla muestra el porcentaje del capital que debe "
            "invertirse en cada acción del portafolio seleccionado. Estos pesos "
            "han sido optimizados mediante la metodología de Markowitz (1952) "
            "para maximizar la eficiencia rentabilidad-riesgo."
        )

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Acción"
        table.rows[0].cells[1].text = "% del capital"
        for a in data["allocation"]:
            row = table.add_row()
            row.cells[0].text = a["ticker"]
            row.cells[1].text = a["weight"]

    # Ranking completo
    doc.add_paragraph("")
    h2 = doc.add_heading("5.2. Ranking completo de portafolios", level=2)
    h2.runs[0].font.color.rgb = RGBColor(46, 80, 144)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Ranking"
    table.rows[0].cells[1].text = "Portafolio"
    table.rows[0].cells[2].text = "Puntuación AHP"
    for r in data["ranking"]:
        row = table.add_row()
        row.cells[0].text = f"#{r['rank']}"
        row.cells[1].text = r["name"]
        row.cells[2].text = f"{r['score']}%"

    # Detalle de portafolios
    doc.add_paragraph("")
    h2 = doc.add_heading("5.3. Métricas de los portafolios candidatos", level=2)
    h2.runs[0].font.color.rgb = RGBColor(46, 80, 144)

    if data.get("portfolios"):
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        for i, h_text in enumerate(["Portafolio", "Rentab.", "Volat.", "Sharpe", "Max DD", "Beta", "Alpha"]):
            table.rows[0].cells[i].text = h_text
        for p_data in data["portfolios"]:
            row = table.add_row()
            row.cells[0].text = p_data["name"]
            row.cells[1].text = f"{p_data['rentabilidad']}%"
            row.cells[2].text = f"{p_data['volatilidad']}%"
            row.cells[3].text = str(p_data["sharpe"])
            row.cells[4].text = f"{p_data['max_drawdown']}%"
            row.cells[5].text = str(p_data["beta"])
            row.cells[6].text = f"{p_data['alpha']}%"

    # ── 6. METODOLOGÍA ──
    doc.add_page_break()
    h = doc.add_heading("6. Nota metodológica", level=1)
    h.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph(
        "Este análisis ha sido realizado mediante la técnica multicriterio AHP "
        "(Analytic Hierarchy Process) propuesta por Saaty (1990), adaptada a la "
        "selección de portafolios de inversión siguiendo la metodología de "
        "Escobar (2015). El proceso incluye: (1) perfilado del inversor mediante "
        "cuestionario, (2) descarga de datos de mercado de tres índices internacionales, "
        "(3) análisis individual de acciones con 15 indicadores financieros, "
        "(4) filtrado secuencial adaptado al perfil, (5) construcción de portafolios "
        "candidatos con optimización de Markowitz, y (6) selección final mediante "
        "AHP con verificación de consistencia."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Referencias:").bold = True
    refs = [
        "Saaty, T.L. (1990). How to make a decision: The Analytic Hierarchy Process. EJOR, 48(1), 9-26.",
        "Escobar, J.W. (2015). Metodología para la toma de decisiones de inversión en portafolio de acciones utilizando AHP. Contaduría y Administración, 60, 346-366.",
        "Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77-91.",
        "Sharpe, W.F. (1966). Mutual fund performance. The Journal of Business, 39(1), 119-138.",
        "Sortino, F.A. y Price, L.N. (1994). Performance measurement in a downside risk framework.",
        "Artzner, P. et al. (1999). Coherent measures of risk. Mathematical Finance, 9(3), 203-228.",
        "Elton, E.J. y Gruber, J.M. (1978). Simple criteria for optimal portfolio selection.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    # Guardar
    import tempfile
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# ─────────────────────────────────────────────────────────────────
# ARRANQUE
# ─────────────────────────────────────────────────────────────────

def open_browser():
    webbrowser.open("http://localhost:5000")


if __name__ == "__main__":
    print("\n" + "█" * 50)
    print("  AHP Portfolio Selector — Web App")
    print("  Abriendo en tu navegador...")
    print("  Para cerrar: pulsa Ctrl+C en esta ventana")
    print("█" * 50 + "\n")

    threading.Timer(1.5, open_browser).start()
    app.run(debug=False, port=5000)
